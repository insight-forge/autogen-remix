from typing import Any, Callable, Dict, List, Literal, Optional, Tuple, Type, Union, Set
import os, json, logging, re, base64
import streamlit
from docx import Document
from io import BytesIO
from PIL import Image
import openai
from mimetypes import guess_type

client_img = openai.AzureOpenAI(
    api_key="",
    azure_endpoint="",
    api_version=""
)


def config_list_from_json(
        env_or_file: str,
        file_location: Optional[str] = ""
) -> List[Dict]:
    json_str = os.environ.get(env_or_file)
    if json_str:
        config_list = json.loads(json_str)
    else:
        config_list_path = os.path.join(file_location, env_or_file)
        try:
            with open(config_list_path) as json_file:
                config_list = json.load(json_file)
        except FileNotFoundError:
            logging.warning(f"The specified config_list file '{config_list_path}' does not exist.")
            return []
    return config_list


def _format_json_str(jstr):
    result = []
    inside_quotes = False
    last_char = " "
    for char in jstr:
        if last_char != "\\" and char == '"':
            inside_quotes = not inside_quotes
        last_char = char
        if not inside_quotes and char == "\n":
            continue
        if inside_quotes and char in ("\n", "\t"):
            char = "\\" + char
        result.append(char)
    return "".join(result)


def local_image_to_data_url(image_path):
    # Guess the MIME type of the image based on the file extension
    mime_type, _ = guess_type(image_path)
    if mime_type is None:
        mime_type = 'application/octet-stream'  # Default MIME type if none is found

    # Read and encode the image file
    with open(image_path, "rb") as image_file:
        base64_encoded_data = base64.b64encode(image_file.read()).decode('utf-8')

    # Construct the data URL
    return f"data:{mime_type};base64,{base64_encoded_data}"


def chat_hisory_process(doc, st):
    text_content = ""
    img_desc = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # 如果段落非空
            text_content += paragraph.text + "\n"

    cnt = 1
    for i in range(len(doc.part.rels)):
        if "image" in doc.part.rels['rId' + str(i + 1)].reltype:
            image_stream = BytesIO(doc.part.rels['rId' + str(i + 1)].target_part.blob)
            img = Image.open(image_stream)
            img_path = 'application/chat_summary/img_tmp/image' + str(cnt) + '.jpg'
            img.save(img_path, 'JPEG')
            data_url = local_image_to_data_url(img_path)

            # 生成图片描述
            response = client_img.chat.completions.create(
                model="gpt-4-vision-preview",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": [
                        {
                            "type": "text",
                            "text": "50字以内描述一下这幅照片"
                        },
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": data_url
                            }
                        }
                    ]}
                ],
                max_tokens=2000
            )
            print(response.choices[0].message.content)
            st.markdown(f"<span style='color:red;'>image{cnt}:{response.choices[0].message.content}</span>",
                        unsafe_allow_html=True)
            # text_content = text_content.replace('{**image' + str(cnt) + '**}', "<*image*>" + response.choices[0].message.content)
            img_desc.append("<*image*>" + response.choices[0].message.content)
            cnt += 1

    for i, paragraph in enumerate(doc.paragraphs):
        for run in paragraph.runs:
            sender = ""
            if "<w:drawing" in run._r.xml:
                try:
                    for j in range(i + 1, len(doc.paragraphs)):
                        if doc.paragraphs[j] and doc.paragraphs[j].runs and doc.paragraphs[j].runs[0].text:
                            sender = doc.paragraphs[j].runs[0].text.split(':')[0]
                            break
                except:
                    pass
                run.text = sender + ": " + img_desc[0]
                img_desc = img_desc[1:]
        if paragraph.text.strip():  # 如果段落非空
            text_content += paragraph.text + "\n"

    return text_content